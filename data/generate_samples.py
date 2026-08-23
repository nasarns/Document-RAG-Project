import os
from pathlib import Path
import docx
import pandas as pd
import pymupdf

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "data" / "samples"
SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

# 1. Sample PDF: Artificial Intelligence & Retrieval Augmented Generation
pdf_path = SAMPLES_DIR / "RAG_Technology_Whitepaper.pdf"
doc = pymupdf.open()

# Page 1
p1 = doc.new_page()
p1.insert_text((50, 60), "Artificial Intelligence & RAG Architecture Whitepaper", fontsize=16)
p1.insert_text((50, 90), "Author: Research Engineering Group\nDate: 2026\nVersion: 1.0", fontsize=10)
p1.insert_text((50, 150), 
    "1. Executive Summary\n"
    "Retrieval-Augmented Generation (RAG) bridges the gap between static LLM memory and dynamic external data sources.\n"
    "By retrieving authoritative context prior to generating answers, hallucination is minimized and factual precision is maximized.\n"
    "This system provides enterprise teams with verified citations and traceability back to source pages.", 
    fontsize=11
)

# Page 2
p2 = doc.new_page()
p2.insert_text((50, 60), "2. Technical Benchmarks & Vector Databases", fontsize=16)
p2.insert_text((50, 100), 
    "The system utilizes Qdrant as the vector database, paired with BGE-small-en embeddings (384 dimensions).\n"
    "Under benchmark testing with 50,000 document chunks:\n"
    "- Retrieval Latency: 12 milliseconds (p95)\n"
    "- Question Answering Accuracy: 96.4%\n"
    "- Cosine similarity threshold for evidence validation: 0.35\n"
    "- Chunking configuration: 700 characters with 150 character sliding overlap.",
    fontsize=11
)
doc.save(str(pdf_path))
doc.close()

# 2. Sample DOCX: Product Specifications
docx_path = SAMPLES_DIR / "Cloud_Platform_Specs.docx"
doc_file = docx.Document()
doc_file.add_heading("Cloud Platform Specifications", level=1)
doc_file.add_paragraph("This document outlines the operational limits, SLA agreements, and security compliance of the Cloud Platform.")
doc_file.add_heading("Section 1: SLA and Uptime", level=2)
doc_file.add_paragraph("The platform guarantees 99.99% monthly uptime. Maintenance windows are scheduled on Sunday at 02:00 UTC.")
doc_file.add_heading("Section 2: Security & Encryption", level=2)
doc_file.add_paragraph("All data at rest is encrypted using AES-256 GCM. Data in transit requires TLS 1.3 encryption.")
doc_file.save(str(docx_path))

# 3. Sample CSV: Product Catalog & Pricing
csv_path = SAMPLES_DIR / "Product_Catalog_2026.csv"
df = pd.DataFrame({
    "Product_ID": ["SKU-1001", "SKU-1002", "SKU-1003", "SKU-1004"],
    "Product_Name": ["UltraBook Pro 16", "Ergonomic Mechanical Keyboard", "27-inch 4K Studio Monitor", "Noise Cancelling Headphones"],
    "Category": ["Computers", "Peripherals", "Displays", "Audio"],
    "Price_USD": [1499.99, 129.50, 449.00, 299.00],
    "Warranty_Years": [3, 2, 3, 1]
})
df.to_csv(csv_path, index=False)

# 4. Sample TXT: Company Policy
txt_path = SAMPLES_DIR / "Remote_Work_Policy.txt"
txt_path.write_text(
    "Global Remote Work Policy 2026\n\n"
    "Eligibility: All full-time engineering and product team members can work remotely up to 5 days per week.\n"
    "Home Office Stipend: Every employee is eligible for an annual home office setup reimbursement of $1,500 USD.\n"
    "Core Collaboration Hours: 10:00 AM to 3:00 PM EST for meetings and standups.",
    encoding="utf-8"
)

print("[DONE] Sample test documents created in data/samples/:")
print(f" - {pdf_path.name}")
print(f" - {docx_path.name}")
print(f" - {csv_path.name}")
print(f" - {txt_path.name}")
