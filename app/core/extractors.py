import os
from pathlib import Path
from typing import List, Dict, Any
import pymupdf  # PyMuPDF
import docx
import pandas as pd


class DocumentExtractionError(Exception):
    """Raised when an error occurs during document text extraction."""
    pass


class DocumentExtractor:
    """
    Extracts readable text and location metadata from various file formats
    (PDF, DOCX, TXT, Excel, CSV) for citation-aware RAG pipelines.
    """

    @classmethod
    def extract(cls, file_path: str | Path) -> List[Dict[str, Any]]:
        """
        Extract text blocks with location metadata from the given file.

        Returns:
            List of dicts: [
                {
                    "text": str,
                    "metadata": {
                        "filename": str,
                        "file_type": str,
                        "page_number": int | None,
                        "sheet_name": str | None,
                        "section": str | None
                    }
                },
                ...
            ]
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        filename = path.name

        if ext == ".pdf":
            return cls._extract_pdf(path, filename)
        elif ext in [".docx", ".doc"]:
            return cls._extract_docx(path, filename)
        elif ext == ".txt":
            return cls._extract_txt(path, filename)
        elif ext in [".xlsx", ".xls"]:
            return cls._extract_excel(path, filename)
        elif ext == ".csv":
            return cls._extract_csv(path, filename)
        else:
            raise DocumentExtractionError(f"Unsupported file format: {ext}")

    @staticmethod
    def _extract_pdf(path: Path, filename: str) -> List[Dict[str, Any]]:
        blocks = []
        try:
            doc = pymupdf.open(str(path))
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                page_text = page.get_text("text").strip()
                if page_text:
                    blocks.append({
                        "text": page_text,
                        "metadata": {
                            "filename": filename,
                            "file_type": "pdf",
                            "page_number": page_idx + 1,
                            "total_pages": len(doc),
                            "sheet_name": None,
                            "section": f"Page {page_idx + 1}"
                        }
                    })
            doc.close()
        except Exception as e:
            raise DocumentExtractionError(f"Error parsing PDF '{filename}': {str(e)}") from e

        if not blocks:
            raise DocumentExtractionError(f"No extractable text found in PDF '{filename}'.")
        return blocks

    @staticmethod
    def _extract_docx(path: Path, filename: str) -> List[Dict[str, Any]]:
        blocks = []
        try:
            doc = docx.Document(str(path))
            current_section = "General"
            current_paras = []

            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if not text:
                    continue

                if p.style and p.style.name.startswith("Heading"):
                    if current_paras:
                        blocks.append({
                            "text": "\n".join(current_paras),
                            "metadata": {
                                "filename": filename,
                                "file_type": "docx",
                                "page_number": None,
                                "sheet_name": None,
                                "section": current_section
                            }
                        })
                        current_paras = []
                    current_section = text
                else:
                    current_paras.append(text)

            if current_paras:
                blocks.append({
                    "text": "\n".join(current_paras),
                    "metadata": {
                        "filename": filename,
                        "file_type": "docx",
                        "page_number": None,
                        "sheet_name": None,
                        "section": current_section
                    }
                })

            # Also extract tables from docx
            for t_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_rows.append(" | ".join(row_data))
                if table_rows:
                    blocks.append({
                        "text": "\n".join(table_rows),
                        "metadata": {
                            "filename": filename,
                            "file_type": "docx",
                            "page_number": None,
                            "sheet_name": None,
                            "section": f"Table {t_idx + 1}"
                        }
                    })
        except Exception as e:
            raise DocumentExtractionError(f"Error parsing DOCX '{filename}': {str(e)}") from e

        if not blocks:
            raise DocumentExtractionError(f"No extractable text found in DOCX '{filename}'.")
        return blocks

    @staticmethod
    def _extract_txt(path: Path, filename: str) -> List[Dict[str, Any]]:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read().strip()
            if not content:
                raise DocumentExtractionError(f"TXT file '{filename}' is empty.")
            return [{
                "text": content,
                "metadata": {
                    "filename": filename,
                    "file_type": "txt",
                    "page_number": 1,
                    "sheet_name": None,
                    "section": "Full Document"
                }
            }]
        except Exception as e:
            raise DocumentExtractionError(f"Error reading TXT '{filename}': {str(e)}") from e

    @staticmethod
    def _extract_excel(path: Path, filename: str) -> List[Dict[str, Any]]:
        blocks = []
        try:
            excel_file = pd.ExcelFile(str(path))
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                if df.empty:
                    continue

                # Convert dataframe into markdown / text chunks in batches of rows
                batch_size = 20
                for start_row in range(0, len(df), batch_size):
                    end_row = min(start_row + batch_size, len(df))
                    chunk_df = df.iloc[start_row:end_row]
                    table_str = chunk_df.to_markdown(index=False)
                    blocks.append({
                        "text": f"Sheet: {sheet_name} (Rows {start_row + 1}-{end_row})\n\n{table_str}",
                        "metadata": {
                            "filename": filename,
                            "file_type": "excel",
                            "page_number": None,
                            "sheet_name": sheet_name,
                            "section": f"Sheet: {sheet_name}, Rows {start_row + 1}-{end_row}"
                        }
                    })
        except Exception as e:
            raise DocumentExtractionError(f"Error parsing Excel '{filename}': {str(e)}") from e

        if not blocks:
            raise DocumentExtractionError(f"No valid data found in Excel '{filename}'.")
        return blocks

    @staticmethod
    def _extract_csv(path: Path, filename: str) -> List[Dict[str, Any]]:
        blocks = []
        try:
            df = pd.read_csv(str(path))
            if df.empty:
                raise DocumentExtractionError(f"CSV file '{filename}' is empty.")

            batch_size = 25
            for start_row in range(0, len(df), batch_size):
                end_row = min(start_row + batch_size, len(df))
                chunk_df = df.iloc[start_row:end_row]
                table_str = chunk_df.to_markdown(index=False)
                blocks.append({
                    "text": f"CSV Records (Rows {start_row + 1}-{end_row})\n\n{table_str}",
                    "metadata": {
                        "filename": filename,
                        "file_type": "csv",
                        "page_number": None,
                        "sheet_name": "CSV",
                        "section": f"Rows {start_row + 1}-{end_row}"
                    }
                })
        except Exception as e:
            raise DocumentExtractionError(f"Error parsing CSV '{filename}': {str(e)}") from e

        return blocks
